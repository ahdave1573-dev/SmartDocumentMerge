"""utils/pdf_to_word.py — Extract text from PDF and save as DOCX"""
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def pdf_to_word(input_path: str, output_path: str) -> str:
    """
    Extract text from each PDF page and write it into a DOCX file.
    Each page is separated by a heading and a horizontal rule.

    Note: This is a text-only extraction; complex layouts (tables, columns)
    are not preserved.
    """
    reader = PdfReader(input_path)
    doc    = Document()

    # Document title
    title = doc.add_heading('Extracted from PDF', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, page in enumerate(reader.pages):
        # Page separator heading
        page_heading = doc.add_heading(f'Page {i + 1}', level=2)
        page_heading.runs[0].font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

        # Extract text
        text = page.extract_text() or ''
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if line:
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(11) if p.runs else None
            else:
                doc.add_paragraph('')  # blank line

        # Add a paragraph rule between pages
        if i < len(reader.pages) - 1:
            doc.add_paragraph('─' * 60)

    doc.save(output_path)
    return output_path
